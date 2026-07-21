from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

ROOT = Path('/Users/mac/Desktop/个人项目/3D工作台')
OUT = ROOT / '3D 导演台操作手册.docx'
SOURCE = Path('/var/folders/q2/tyst5rxs6fbf147j82tf8knr0000gn/T/codex-clipboard-d7b87ec3-32b1-41a0-b893-6a972022c679.png')
ASSET = ROOT / 'tmp' / 'manual-timeline.png'

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_run_font(run, name='Microsoft YaHei', size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(doc, text='', style=None, before=0, after=5, line=1.15):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if text:
        r = p.add_run(text)
        set_run_font(r)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r)
    return p

def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run_font(r, size=15 if level == 1 else 12, bold=True, color=(31, 78, 121) if level == 1 else (45, 45, 45))
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9, color=(100, 100, 100))
    return p

def add_image(doc, path, width=6.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    return p

def add_tip(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F2F5F8')
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + '  ')
    set_run_font(r, size=10, bold=True, color=(31, 78, 121))
    r = p.add_run(text)
    set_run_font(r, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def main():
    ASSET.parent.mkdir(parents=True, exist_ok=True)
    if SOURCE.exists():
        with Image.open(SOURCE) as im:
            # Keep the whole demo for design mode and crop the timeline for animation mode.
            im.crop((250, 600, 1655, 905)).save(ASSET)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    normal = doc.styles['Normal']
    normal.font.name = 'Microsoft YaHei'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run('3D 导演台操作手册')
    set_run_font(r, size=20, bold=True, color=(31, 78, 121))
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('面向运营同学｜场景搭建、动画制作与视频导出')
    set_run_font(r, size=10, color=(100, 100, 100))
    add_para(doc, '本手册用于快速完成一次从搭建场景到导出视频的操作。界面以当前 demo 为准。', after=10)

    add_heading(doc, '一、开始使用', 1)
    add_number(doc, '进入导演台后，默认处于“设计”模式。')
    add_number(doc, '左侧选择机位或对象，中间视窗进行观察和调整，右侧查看并修改属性。')
    add_number(doc, '需要制作动画时，切换底部“动画”模式，打开时间轴。')
    add_tip(doc, '模式区别', '设计模式用于搭建和调整场景；动画模式用于记录关键帧、安排机位切换和导出视频。')

    add_heading(doc, '二、设计模式', 1)
    add_heading(doc, '1. 添加与管理资产', 2)
    add_bullet(doc, '左侧“资产列表”按“机位、对象”查看当前场景。')
    add_bullet(doc, '通过导入或添加功能创建资产；点击资产可在右侧查看属性。')
    add_bullet(doc, '眼睛图标控制显示 / 隐藏，锁定图标控制是否允许编辑，删除图标移除资产。')
    add_bullet(doc, '隐藏只影响画面展示和播放，资产仍可编辑；删除后会同步移除其时间轴内容。')
    add_heading(doc, '2. 调整视窗与资产', 2)
    add_bullet(doc, '在中间视窗拖动、旋转或缩放工具，调整资产的空间状态。')
    add_bullet(doc, '右侧属性面板可修改名称、位置、旋转、缩放等参数；角色还可切换姿势预设并展开身体、手臂、腿部和头部调节。')
    add_bullet(doc, '机位可调整位置、朝向、焦距，并选择自由朝向或注视目标模式。')
    add_heading(doc, '3. 画面管理', 2)
    add_bullet(doc, '右侧“画面管理”分为“快照”和“视频”两个 Tab。')
    add_bullet(doc, '快照用于保存当前画面；视频用于保存录制结果，后续在视频 Tab 中选择并导出。')
    add_heading(doc, '4. 固定机位视角', 2)
    add_bullet(doc, '在左侧机位操作中点击“固定机位视角”，确认后进入该机位视角。')
    add_bullet(doc, '固定后视角不可通过鼠标或快捷键移动；关闭提示或退出固定视角后恢复自由视角。')
    if SOURCE.exists():
        add_image(doc, SOURCE, 6.5)
        add_caption(doc, '图 1 设计模式：左侧资产列表、中间视窗、右侧属性与画面管理')

    add_heading(doc, '三、动画模式', 1)
    add_heading(doc, '1. 打开时间轴', 2)
    add_bullet(doc, '点击底部“动画”切换到动画模式，时间轴显示在视窗下方。')
    add_bullet(doc, '时间轴为空时，只显示引导文案；先在左侧选中资产，再点击“插帧”添加。')
    add_bullet(doc, '时间轴中的资产与左侧资产对应，但关键帧记录的是各时间点的参数快照。')
    add_heading(doc, '2. 添加轨道与关键帧', 2)
    add_number(doc, '在左侧选中机位或对象。')
    add_number(doc, '把 CTI 播放头移动到目标时间。')
    add_number(doc, '点击“插帧”，生成对应轨道，并在当前 CTI 位置生成默认关键帧。')
    add_number(doc, '调整资产参数后，点击“插帧”记录新的状态；开启自动插帧时，属性修改会自动记录。')
    add_heading(doc, '3. 编辑关键帧', 2)
    add_bullet(doc, '拖动菱形关键帧可改变发生时间；点击关键帧可查看和编辑该时刻参数。')
    add_bullet(doc, '删除按钮用于删除选中的关键帧或时间轴轨道。')
    add_bullet(doc, '对象关键帧记录位置、旋转、缩放、统一缩放及角色姿势；姿势在关键帧处瞬间切换，不做连续插值。')
    add_bullet(doc, '机位关键帧记录位置、旋转 / 朝向、注视目标或坐标、焦段等相机参数。')
    add_heading(doc, '4. 设置机位序列', 2)
    add_number(doc, '点击“机位序列”右侧加号，从左侧资产列表选择要加入的机位。')
    add_number(doc, '在时间轴上添加机位帧；第一个机位帧生效到下一个机位帧，之后依次切换。')
    add_number(doc, '拖动机位帧可调整切换时间；同一序列中的机位帧不重叠。')
    add_bullet(doc, '播放视角优先级：机位序列 > 固定机位视角 > 自由视角。')
    add_heading(doc, '5. 播放、范围与导出', 2)
    add_bullet(doc, '拖动 CTI 可实时查看当前时间点的画面；播放时会读取关键帧之间的状态。')
    add_bullet(doc, '设置入点和出点，确定最终视频的录制范围；未设置时按时间轴默认范围导出。')
    add_bullet(doc, '点击“导出视频”开始录制。录制过程中按提示等待完成；中途取消则本次视频不保存。')
    add_bullet(doc, '完成后到右侧“画面管理 > 视频”中选择视频，再进行导出。')
    if ASSET.exists():
        add_image(doc, ASSET, 6.5)
        add_caption(doc, '图 2 动画模式：时间轴、机位序列、轨道与关键帧')

    add_heading(doc, '四、常用规则', 1)
    add_bullet(doc, '点击左侧资产：右侧显示资产属性，并取消时间轴选中状态。')
    add_bullet(doc, '点击时间轴轨道或关键帧：取消左侧资产选中，画布切换到对应时间点状态。其他轨道保持正常显示。')
    add_bullet(doc, '设计模式与动画模式的撤销 / 重做记录相互隔离。Mac 使用 Command + Z 撤销、Command + Shift + Z 重做；Windows 使用 Ctrl + Z、Ctrl + Shift + Z。')
    add_bullet(doc, '快捷键：V 移动，R 旋转，X 缩放。')
    add_bullet(doc, '项目会实时保存。')

    add_heading(doc, '五、快速操作清单', 1)
    add_number(doc, '设计模式：添加机位和对象，完成位置、姿势、视角调整。')
    add_number(doc, '动画模式：选中资产，移动 CTI，点击“插帧”。')
    add_number(doc, '继续调整参数并添加关键帧，必要时安排机位序列。')
    add_number(doc, '设置入点 / 出点，播放检查，点击“导出视频”。')
    add_number(doc, '到“画面管理 > 视频”中选择并导出成片。')

    doc.save(OUT)
    print(OUT)

if __name__ == '__main__':
    main()
